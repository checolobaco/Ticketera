import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_standalone_mandate_docx(filename):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Base Styles
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x27, 0x27, 0x2A)  # #27272a
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Palette
    NAVY = RGBColor(0x0F, 0x17, 0x2A)       # #0f172a
    DARK_BLUE = RGBColor(0x1E, 0x3A, 0x8A)  # #1e3a8a

    # Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CONTRATO DE MANDATO EXCLUSIVO PARA EL RECAUDO DE FONDOS VÍA WOMPI")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = NAVY
    title_p.paragraph_format.space_after = Pt(4)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("(RECAUDO POR CUENTA DE TERCEROS - ARTÍCULO 1262 DEL CÓDIGO DE COMERCIO DE COLOMBIA)")
    sub_run.bold = True
    sub_run.font.size = Pt(10.5)
    sub_run.font.color.rgb = DARK_BLUE
    subtitle_p.paragraph_format.space_after = Pt(16)

    # Intro / Preamble
    intro_p = doc.add_paragraph()
    intro_p.add_run("Entre los suscritos a saber: Por una parte, ")
    intro_p.add_run("[NOMBRE / RAZÓN SOCIAL DEL ORGANIZADOR]").bold = True
    intro_p.add_run(", persona jurídica/natural identificada con NIT / C.C. No. ")
    intro_p.add_run("[NIT O CÉDULA DEL MANDANTE]").bold = True
    intro_p.add_run(", domiciliada en la ciudad de [CIUDAD], representada legalmente por ")
    intro_p.add_run("[NOMBRE REPRESENTANTE LEGAL]").bold = True
    intro_p.add_run(", identificado(a) con C.C. No. [NÚMERO CÉDULA], en adelante denominado ")
    intro_p.add_run("EL MANDANTE").bold = True
    intro_p.add_run("; y por la otra parte, ")
    intro_p.add_run("[NOMBRE / RAZÓN SOCIAL DEL TITULAR DE WOMPI]").bold = True
    intro_p.add_run(", persona jurídica/natural identificada con NIT / C.C. No. ")
    intro_p.add_run("[NIT/C.C. DEL MANDATARIO]").bold = True
    intro_p.add_run(", domiciliada en [CIUDAD], titular de la cuenta de pasarela de pagos Wompi Colombia, en adelante denominado ")
    intro_p.add_run("EL MANDATARIO").bold = True
    intro_p.add_run("; se conviene celebrar el presente ",)
    intro_p.add_run("CONTRATO DE MANDATO EXCLUSIVO DE RECAUDO DE FONDOS VÍA PASARELA WOMPI").bold = True
    intro_p.add_run(", previo conocimiento de las siguientes CONSIDERACIONES:")

    # CONSIDERACIONES
    p_cons = doc.add_paragraph()
    p_cons.paragraph_format.left_indent = Inches(0.2)
    p_cons.paragraph_format.space_after = Pt(4)
    p_cons.add_run("1. ").bold = True
    p_cons.add_run("Que las partes suscribieron o mantienen vigente un acuerdo independiente sobre la provisión de software y soporte para la emisión de boletería digital mediante la plataforma CloudTickets (Contratación de Servicios Tecnológicos).\n")
    p_cons.add_run("2. ").bold = True
    p_cons.add_run("Que para la comercialización en línea de las entradas del evento, las partes acuerdan formalizar en este documento la figura jurídica de MANDATO DE RECAUDO POR CUENTA DE TERCEROS, a fin de regular de manera autónoma la recepción de dinero a través de la pasarela Wompi del MANDATARIO, sus costos financieros, impuestos y liquidaciones.")

    def add_clause_heading(number_title, text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        run1 = h.add_run(number_title)
        run1.bold = True
        run1.font.color.rgb = DARK_BLUE
        run1.font.size = Pt(11.5)
        run2 = h.add_run(" " + text)
        run2.bold = True
        run2.font.color.rgb = NAVY
        run2.font.size = Pt(11.5)
        return h

    # CLAUSULA PRIMERA
    add_clause_heading("CLÁUSULA PRIMERA –", "OBJETO DEL MANDATO DE RECAUDO:")
    doc.add_paragraph(
        "EL MANDANTE encomienda de manera expresa e irrevocable al MANDATARIO, y este acepta, la gestión exclusiva de recaudo por cuenta y riesgo del MANDANTE de los dineros provenientes de las compras de boletería realizadas a través de la pasarela de pagos Wompi Colombia para el evento denominado:"
    )
    
    event_box = doc.add_paragraph()
    event_box.paragraph_format.left_indent = Inches(0.4)
    event_run = event_box.add_run(
        "• Evento: [NOMBRE DEL EVENTO - HARVY VALENCIA TULUÁ]\n"
        "• Fecha y Hora del Evento: [FECHA Y HORA]\n"
        "• Lugar / Sede: [LUGAR DEL EVENTO / CIUDAD]\n"
        "• Dominio / Canal de Venta: [cloud-tickets.com / SUBDOMINIO]"
    )
    event_run.bold = True
    event_run.font.color.rgb = DARK_BLUE

    # CLAUSULA SEGUNDA
    add_clause_heading("CLÁUSULA SEGUNDA –", "NATURALEZA JURÍDICA Y EFECTOS TRIBUTARIOS ANTE LA DIAN:")
    doc.add_paragraph(
        "De conformidad con el Artículo 1262 y subsiguientes del Código de Comercio de Colombia, las partes declaran expresamente:"
    )
    
    points_c2 = [
        ("A. Dominio de los Fondos: ", "Los recursos brutos recaudados mediante la pasarela Wompi por concepto de venta de tickets pertenecen desde su acreditación única y exclusivamente al MANDANTE."),
        ("B. Exoneración de Ingreso Propio: ", "Los dineros ingresados a la cuenta de pasarela o cuenta bancaria del MANDATARIO por concepto de ventas del evento NO constituyen venta propia, ingreso bruto gravable ni incremento patrimonial para EL MANDATARIO ante la Dirección de Impuestos y Aduanas Nacionales (DIAN) ni autoridades locales."),
        ("C. Carácter de Intermediario: ", "EL MANDATARIO actúa en calidad de simple intermediario recaudador por cuenta de terceros. Los únicos ingresos gravables para el MANDATARIO corresponden a las tarifas o comisiones por servicio de intermediación pactadas en el contrato tecnológico.")
    ]
    for title, desc in points_c2:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(title).bold = True
        p.add_run(desc)

    # CLAUSULA TERCERA
    add_clause_heading("CLÁUSULA TERCERA –", "DEDUCCIONES OPERATIVAS Y GASTOS FINANCIEROS AUTORIZADOS:")
    doc.add_paragraph(
        "EL MANDANTE autoriza de manera expresa al MANDATARIO a descontar directamente del monto bruto total recaudado en Wompi, antes de efectuar cualquier giro o transferencia de saldos, los siguientes valores:"
    )

    discounts = [
        ("1. Comisión de Pasarela Wompi: ", "La tarifa cobrada directamente por Wompi Colombia (correspondiente al 2.9% + $590 COP + IVA por transacción aprobada o tarifa de comercio vigente)."),
        ("2. Retenciones Financieras de Ley: ", "Las retenciones en la fuente (ReteFuente, ReteIVA, ReteICA) practicadas por Wompi y el sistema financiero sobre las transacciones del evento."),
        ("3. Gravamen a los Movimientos Financieros (4x1000): ", "El costo del Impuesto al 4x1000 (GMF) que se genere en la cuenta bancaria del MANDATARIO al momento de realizar la transferencia de los fondos hacia la cuenta del MANDANTE."),
        ("4. Comisiones de Servicio / Plataforma: ", "Las comisiones o porcentajes por boletería vendida estipuladas a favor de la plataforma CloudTickets según la propuesta comercial / contrato técnico acordado.")
    ]
    for title, desc in discounts:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(title).bold = True
        p.add_run(desc)

    # CLAUSULA CUARTA
    add_clause_heading("CLÁUSULA CUARTA –", "CRONOGRAMA DE LIQUIDACIÓN Y RESERVA DE GARANTÍA:")
    doc.add_paragraph(
        "El desembolso de los recursos recaudados a favor del MANDANTE se sujetará a las siguientes reglas de liquidación y retención temporal de seguridad:"
    )

    p_l1 = doc.add_paragraph()
    p_l1.paragraph_format.left_indent = Inches(0.2)
    p_l1.add_run("A. Desembolsos Parciales (Anticipos): ").bold = True
    p_l1.add_run("EL MANDATARIO realizará giros del [PORCENTAJE GIRO, ej. 80%] del recaudo neto acumulado en Wompi con cortes [SEMANALES / PREVIOS AL EVENTO].")

    p_l2 = doc.add_paragraph()
    p_l2.paragraph_format.left_indent = Inches(0.2)
    p_l2.add_run("B. Fondo de Reserva de Garantía por Devoluciones: ").bold = True
    p_l2.add_run("EL MANDATARIO retendrá un [PORCENTAJE RESERVA, ej. 15% a 20%] del valor neto total recaudado hasta la realización del evento. Dicha reserva tiene como fin exclusivo respaldar eventuales reclamos, solicitudes de retracto de compra o contracargos bancarios presentados por los compradores de conformidad con el Estatuto del Consumidor (Ley 1480 de 2011).")

    p_l3 = doc.add_paragraph()
    p_l3.paragraph_format.left_indent = Inches(0.2)
    p_l3.add_run("C. Liquidación Definitiva: ").bold = True
    p_l3.add_run("Transcurridos tres a cinco (3 a 5) días hábiles posteriores a la finalización del evento, EL MANDATARIO enviará el reporte de liquidación final consolidando ventas, deducciones de Wompi, 4x1000 y comisiones, y transferirá el saldo restante al MANDANTE.")

    # CLAUSULA QUINTA
    add_clause_heading("CLÁUSULA QUINTA –", "RESPONSABILIDAD POR CANCELACIÓN O DEVOLUCIÓN DE ENTRADAS:")
    doc.add_paragraph(
        "EL MANDANTE asume la responsabilidad exclusiva, total e ilimitada frente a los compradores de entradas y entidades reguladoras por cualquier reclamación originada por cancelación, aplazamiento, modificación de artistas, aforo o fallas en la realización del evento."
    )
    doc.add_paragraph(
        "PARÁGRAFO: Si por causa de cancelación o decisión del MANDANTE se debe reembolsar el dinero a los compradores, EL MANDATARIO efectuará las devoluciones utilizando la reserva retenida o el saldo recaudado en Wompi. En caso de que el saldo sea insuficiente, EL MANDANTE se obliga a consignar la diferencia faltante al MANDATARIO dentro de las 48 horas siguientes a la solicitud formal."
    )

    # CLAUSULA SEXTA
    add_clause_heading("CLÁUSULA SEXTA –", "FACTURACIÓN E INFORMACIÓN TRIBUTARIA:")
    doc.add_paragraph(
        "Para todos los efectos fiscales ante la DIAN:"
    )
    doc.add_paragraph().add_run("• EL MANDANTE ").bold = True
    doc.paragraphs[-1].add_run("es el único obligado a expedir la factura electrónica o documento equivalente de boletería a los asistentes por el valor total de las entradas cobradas a través de Wompi.")
    
    doc.add_paragraph().add_run("• EL MANDATARIO ").bold = True
    doc.paragraphs[-1].add_run("expedirá factura únicamente por las comisiones y honorarios de servicio prestados.")

    # CLAUSULA SEPTIMA
    add_clause_heading("CLÁUSULA SÉPTIMA –", "CUENTA BANCARIA DE DESTINO:")
    doc.add_paragraph(
        "Los desembolsos netos se girarán únicamente a la cuenta bancaria certificada del MANDANTE:"
    )

    bank_box = doc.add_paragraph()
    bank_box.paragraph_format.left_indent = Inches(0.4)
    bank_run = bank_box.add_run(
        "• Banco: [NOMBRE DEL BANCO]\n"
        "• Tipo de Cuenta: [AHORROS / CORRIENTE]\n"
        "• Número de Cuenta: [NÚMERO DE CUENTA]\n"
        "• Titular: [NOMBRE / RAZÓN SOCIAL DEL MANDANTE]\n"
        "• NIT / C.C.: [NIT O CÉDULA DEL TITULAR]"
    )
    bank_run.bold = True
    bank_run.font.color.rgb = DARK_BLUE

    # CLAUSULA OCTAVA
    add_clause_heading("CLÁUSULA OCTAVA –", "FIRMAS Y VALIDEZ ELECTRÓNICA:")
    doc.add_paragraph(
        "De conformidad con la Ley 527 de 1999 de la República de Colombia, el intercambio de este documento firmado de forma digital, manuscrita escaneada o vía correo electrónico constituye prueba fehaciente de la voluntad de las partes."
    )

    doc.add_paragraph(
        "En constancia, se firma en dos (2) ejemplares de igual valor legal en [CIUDAD], a los [DÍA] días del mes de [MES] del año 2026."
    )

    # Signatures
    doc.add_paragraph().paragraph_format.space_before = Pt(24)
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    cell_mandante = sig_table.rows[0].cells[0]
    cell_mandatario = sig_table.rows[0].cells[1]
    
    p_m = cell_mandante.paragraphs[0]
    p_m.add_run("_________________________________________\n").bold = True
    p_m.add_run("EL MANDANTE (ORGANIZADOR DEL EVENTO):\n").bold = True
    p_m.add_run("Firma: _________________________________\n"
              "Nombre: [NOMBRE REPRESENTANTE]\n"
              "C.C. / NIT: [NÚMERO DOC]\n"
              "Evento / Empresa: [HARVY VALENCIA TULUÁ]")

    p_a = cell_mandatario.paragraphs[0]
    p_a.add_run("_________________________________________\n").bold = True
    p_a.add_run("EL MANDATARIO (TITULAR DE WOMPI):\n").bold = True
    p_a.add_run("Firma: _________________________________\n"
              "Nombre: [NOMBRE PROVEEDOR]\n"
              "C.C. / NIT: [NÚMERO DOC]\n"
              "Plataforma: CloudTickets")

    # Remove borders from signature table
    for row in sig_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders {} >
                    <w:top w:val="none"/>
                    <w:left w:val="none"/>
                    <w:bottom w:val="none"/>
                    <w:right w:val="none"/>
                </w:tcBorders>
            '''.format(nsdecls('w')))
            tcPr.append(tcBorders)

    doc.save(filename)
    print(f"Standalone mandate DOCX saved successfully at: {filename}")

if __name__ == "__main__":
    create_standalone_mandate_docx(r"C:\0DE\Ticketera\Clients\Harvi_valencia_tulua\CONTRATO_MANDATO_RECAUDO_WOMPI_HARVY_VALENCIA.docx")
    create_standalone_mandate_docx(r"c:\0DE\Ticketera\CONTRATO_MANDATO_RECAUDO_WOMPI.docx")
