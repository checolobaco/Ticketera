import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_contract_docx(filename):
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x27, 0x27, 0x2A)  # #27272a
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Colors
    NAVY = RGBColor(0x0F, 0x17, 0x2A)       # #0f172a
    DARK_BLUE = RGBColor(0x1E, 0x3A, 0x8A)  # #1e3a8a
    GRAY_TEXT = RGBColor(0x47, 0x55, 0x69)  # #475569

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CONTRATO DE PRESTACIÓN DE SERVICIOS TECNOLÓGICOS\nY LICENCIAMIENTO DE SOFTWARE (CLOUD TICKETS)")
    title_run.bold = True
    title_run.font.size = Pt(15)
    title_run.font.color.rgb = NAVY
    title_p.paragraph_format.space_after = Pt(4)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("VALOR TOTAL DEL CONTRATO: $800.000 COP")
    sub_run.bold = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = DARK_BLUE
    subtitle_p.paragraph_format.space_after = Pt(18)

    # Intro / Preamble
    intro_p = doc.add_paragraph()
    intro_p.add_run(
        "Entre los suscritos a saber: Por una parte, "
    )
    r_client = intro_p.add_run("[NOMBRE / RAZÓN SOCIAL DEL CLIENTE]")
    r_client.bold = True
    intro_p.add_run(
        ", sociedad identificada con NIT / Cédula No. [NIT O CÉDULA CLIENTE], domiciliada en la ciudad de [CIUDAD], representada legalmente por "
    )
    r_rep = intro_p.add_run("[NOMBRE REPRESENTANTE LEGAL CLIENTE]")
    r_rep.bold = True
    intro_p.add_run(
        ", identificado(a) con C.C. No. [NÚMERO CÉDULA], quien en adelante y para los efectos del presente contrato se denominará el "
    )
    r_con = intro_p.add_run("EL CONTRATANTE")
    r_con.bold = True
    intro_p.add_run(
        "; y por la otra parte, "
    )
    r_prov = intro_p.add_run("[NOMBRE / RAZÓN SOCIAL PROVEEDOR CLOUDTICKETS]")
    r_prov.bold = True
    intro_p.add_run(
        ", persona jurídica/natural identificada con NIT / C.C. No. [NIT/C.C. PROVEEDOR], domiciliada en [CIUDAD], quien en adelante se denominará "
    )
    r_prov_title = intro_p.add_run("EL CONTRATISTA")
    r_prov_title.bold = True
    intro_p.add_run(
        "; se ha convenido celebrar el presente ",
    )
    r_doc = intro_p.add_run("CONTRATO DE PRESTACIÓN DE SERVICIOS Y LICENCIAMIENTO DE SOFTWARE")
    r_doc.bold = True
    intro_p.add_run(
        ", el cual se regirá por las siguientes cláusulas y en lo no previsto en ellas por las leyes comerciales y civiles aplicables de la República de Colombia:"
    )

    def add_clause_heading(number_title, text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        run1 = h.add_run(number_title)
        run1.bold = True
        run1.font.color.rgb = DARK_BLUE
        run1.font.size = Pt(11.5)
        run2 = h.add_run(" " + text)
        run2.bold = True
        run2.font.color.rgb = NAVY
        run2.font.size = Pt(11.5)
        return h

    # CLAUSULA PRIMERA
    add_clause_heading("CLÁUSULA PRIMERA –", "OBJETO DEL CONTRATO:")
    p = doc.add_paragraph(
        "EL CONTRATISTA se obliga para con EL CONTRATANTE a prestar los servicios tecnológicos de despliegue, configuración, mantenimiento y licenciamiento de la plataforma digital CLOUDTICKETS (Ticketera), destinada a la comercialización, emisión, entrega automatizada y control de acceso de boletería para el evento denominado:"
    )
    
    event_box = doc.add_paragraph()
    event_box.paragraph_format.left_indent = Inches(0.4)
    event_run = event_box.add_run("• Evento: [NOMBRE DEL EVENTO]\n• Fecha del Evento: [FECHA Y HORA DEL EVENTO]\n• Lugar / Sede: [LUGAR DEL EVENTO / CIUDAD]\n• Aforo Proyectado: [CANTIDAD DE ASISTENTES ESTIMADA]")
    event_run.bold = True
    event_run.font.color.rgb = DARK_BLUE

    # CLAUSULA SEGUNDA
    add_clause_heading("CLÁUSULA SEGUNDA –", "ALCANCE DE LOS SERVICIOS Y ENTREGABLES INCLUIDOS:")
    doc.add_paragraph(
        "Por el valor acordado en el presente contrato, EL CONTRATISTA entregará y pondrá a disposición del CONTRATANTE los siguientes componentes tecnológicos:"
    )

    deliverables = [
        ("A. Portal Web de Venta y Checkout Digital:", " Despliegue de la aplicación web de venta de entradas con catálogo de boletas, selección de localidades y formulario de registro del comprador en dominio personalizado o subdominio asignado."),
        ("B. Pasarela de Pagos & Módulo de Transferencias:", " Integración con la pasarela de pagos Wompi Colombia (para pagos con Tarjeta de Crédito/Débito, PSE, Nequi y Bancolombia) y módulo administrativo para validación de compras por transferencia bancaria con carga de comprobante."),
        ("C. Notificaciones y Entrega Automatizada por WhatsApp y Correo:", " Envío instantáneo del ticket digital en formato PDF de alta definición con código QR dinámico directo a la cuenta de WhatsApp del comprador mediante Meta WhatsApp Cloud API (costo del servicio de mensajería asumido por EL CONTRATISTA), con respaldo transaccional por correo electrónico."),
        ("D. App Móvil de Control de Acceso (QR & NFC):", " Licencia y habilitación de la aplicación móvil de lectura para el personal de puerta del CONTRATANTE, con capacidad de validación por cámara de códigos QR y lectura de tecnología NFC (pulseras/tarjetas para zona VIP o consumo)."),
        ("E. Módulo Administrativo y Analítica:", " Acceso al panel de administración para gestión de inventario de boletería, creación de códigos de descuento promocionales, asignación de beneficios de cortesía (consumibles) y descarga de reportes financieros y de aforo en tiempo real."),
        ("F. Soporte Técnico:", " Asistencia técnica remota y monitoreo de infraestructura durante la fase de venta y el día del evento para resolver incidencias de acceso en puerta.")
    ]

    for title, desc in deliverables:
        p_del = doc.add_paragraph()
        p_del.paragraph_format.left_indent = Inches(0.2)
        p_del.paragraph_format.space_after = Pt(4)
        r_t = p_del.add_run(title)
        r_t.bold = True
        r_t.font.color.rgb = DARK_BLUE
        p_del.add_run(desc)

    # CLAUSULA TERCERA
    add_clause_heading("CLÁUSULA TERCERA –", "VALOR DEL CONTRATO Y FORMA DE PAGO:")
    p_val = doc.add_paragraph(
        "El valor total pactado por la prestación de los servicios tecnológicos detallados en la Cláusula Segunda es de "
    )
    r_v = p_val.add_run("OCHOCIENTOS MIL PESOS COLOMBIANOS M/CTE ($800.000 COP)")
    r_v.bold = True
    p_val.add_run(", los cuales serán cancelados por EL CONTRATANTE según el siguiente esquema de pagos:")

    # Table for Payment Schedule
    table = doc.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Hito de Pago", "Porcentaje / Monto", "Condición de Pago"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        shading = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    data = [
        ("Anticipo de Despliegue", "50%  ($400.000 COP)", "A la firma del contrato como requisito para inicio de configuración y despliegue cloud."),
        ("Saldo Final (Go-Live)", "50%  ($400.000 COP)", "A la habilitación del sitio para salida a ventas o máximo 24 horas antes del evento.")
    ]

    for row_idx, row_data in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = text
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for r in p.runs:
                r.font.size = Pt(9.5)
            if col_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.bold = True
            # Alternating shading
            if row_idx % 2 == 1:
                shading = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
                row_cells[col_idx]._tc.get_or_add_tcPr().append(shading)

    # Bank details paragraph
    p_bank = doc.add_paragraph()
    p_bank.paragraph_format.space_before = Pt(8)
    p_bank.add_run("PARÁGRAFO PRIMERO (Medio de Pago): ").bold = True
    p_bank.add_run("Los pagos deberán realizarse mediante transferencia bancaria a la cuenta [AHORROS/CORRIENTE] No. [NÚMERO CUENTA] del banco [NOMBRE BANCO] a nombre de [TITULAR CUENTA] identificada con NIT/C.C. [NÚMERO IDENTIFICACIÓN]. EL CONTRATISTA emitirá la correspondiente Factura de Venta o Cuenta de Cobro por cada pago recibido.")

    # CLAUSULA CUARTA (MODIFICADA SEGÚN INDICACIONES DEL USUARIO)
    add_clause_heading("CLÁUSULA CUARTA –", "EXCLUSIONES, HARDWARE Y CONDICIONES DE SERVICIOS:")
    doc.add_paragraph(
        "Las partes dejan constancia explícita y de común acuerdo sobre las siguientes condiciones operativas y de costos asociados al servicio:"
    )

    exclusions = [
        ("1. No Suministro de Dispositivos Electrónicos ni de Cómputo: ", "EL CONTRATISTA no suministra, no provee ni alquila ningún tipo de dispositivo electrónico o equipo físico de cómputo, tales como celulares, teléfonos inteligentes, tablets, computadores portátiles o de escritorio, impresoras de tickets, ni lectores físicos independientes de código de barras o tarjetas NFC. Corresponde exclusivamente al CONTRATANTE disponer de los dispositivos móviles inteligentes (Android / iOS) con cámara integrada y/o lector NFC, así como de la logística y personal necesario para la lectura de boletas en puerta."),
        ("2. Costos de Notificaciones por WhatsApp (Asumidos por el Contratista): ", "Los costos de mensajería y notificaciones automatizadas generados a través de la API oficial de Meta WhatsApp Cloud API para el envío de las entradas digitales en formato PDF a los compradores serán totalmente asumidos e incluidos por EL CONTRATISTA dentro del valor global de $800.000 COP contratado, sin generar ningún cobro adicional o posterior para EL CONTRATANTE."),
        ("3. Comisiones de Pasarela de Pagos (Wompi): ", "La tarifa correspondiente a la pasarela de pagos Wompi Colombia (2.9% + $590 COP + IVA por transacción exitosa o tarifa contratada) es deducida directamente por la pasarela sobre los fondos recaudados o facturada de forma independiente al CONTRATANTE.")
    ]

    for title, desc in exclusions:
        p_exc = doc.add_paragraph()
        p_exc.paragraph_format.left_indent = Inches(0.2)
        p_exc.paragraph_format.space_after = Pt(4)
        p_exc.add_run(title).bold = True
        p_exc.add_run(desc)

    # CLAUSULA QUINTA
    add_clause_heading("CLÁUSULA QUINTA –", "OBLIGACIONES DE LAS PARTES:")
    
    doc.add_paragraph().add_run("A. Obligaciones del CONTRATISTA:").bold = True
    ob_prov = [
        "Desplegar y mantener operativa la infraestructura en la nube para la venta de tickets durante la vigencia del contrato.",
        "Garantizar la correcta generación y envío de las entradas con código QR único a través de WhatsApp (asumiendo sus costos de mensajería) y correo electrónico.",
        "Suministrar el acceso al panel administrativo y capacitar al personal del CONTRATANTE en el uso de la App Móvil Lectora.",
        "Brindar soporte técnico remoto durante la jornada del evento para resolver inconsistencias o incidentes de lectura."
    ]
    for item in ob_prov:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.left_indent = Inches(0.3)
        p_item.paragraph_format.space_after = Pt(2)
        p_item.add_run("• " + item)

    doc.add_paragraph().add_run("B. Obligaciones del CONTRATANTE:").bold = True
    ob_client = [
        "Suministrar oportunamente la información del evento (logos, precios, términos, aforos y políticas de reembolso).",
        "Disponer y proveer de los dispositivos móviles inteligentes (smartphones/tablets) y la conectividad a internet adecuada en el recinto para la validación en puerta por parte de su personal.",
        "Efectuar los pagos pactados en los tiempos y condiciones establecidas en la Cláusula Tercera.",
        "Asumir la total responsabilidad civil, comercial y legal ante los compradores por la realización, reprogramación o cancelación del evento."
    ]
    for item in ob_client:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.left_indent = Inches(0.3)
        p_item.paragraph_format.space_after = Pt(2)
        p_item.add_run("• " + item)

    # CLAUSULA SEXTA
    add_clause_heading("CLÁUSULA SEXTA –", "PROPIEDAD INTELECTUAL Y LICENCIAMIENTO:")
    doc.add_paragraph(
        "EL CONTRATISTA otorga al CONTRATANTE una licencia de uso no exclusiva, temporal y revocable sobre la plataforma CLOUDTICKETS únicamente para la realización del evento contratado. Todos los derechos de propiedad intelectual, marcas, marcas registradas, código fuente, arquitecturas de software y desarrollos a medida pertenecen y seguirán perteneciendo de forma exclusiva al CONTRATISTA."
    )

    # CLAUSULA SEPTIMA
    add_clause_heading("CLÁUSULA SÉPTIMA –", "PROTECCIÓN DE DATOS PERSONALES (HABEAS DATA):")
    doc.add_paragraph(
        "En cumplimiento de la Ley 1581 de 2012 y el Decreto 1377 de 2013 de la República de Colombia, ambas partes se comprometen a dar un tratamiento adecuado y seguro a los datos personales recolectados de los compradores de boletería. EL CONTRATISTA actuará como Encargado del Tratamiento y aplicará las medidas de seguridad necesarias para salvaguardar la confidencialidad de la información."
    )

    # CLAUSULA OCTAVA
    add_clause_heading("CLÁUSULA OCTAVA –", "GARANTÍA Y LIMITACIÓN DE RESPONSABILIDAD:")
    doc.add_paragraph(
        "EL CONTRATISTA garantiza el funcionamiento de la plataforma en condiciones normales de uso. No obstante, no será responsable por fallas derivadas de caídas generales de conectividad a internet en el recinto del evento, fallas externas en los servidores de la pasarela Wompi o Meta WhatsApp, ni por fuerza mayor o caso fortuito debidamente comprobados."
    )

    # CLAUSULA NOVENA
    add_clause_heading("CLÁUSULA NOVENA –", "DURACIÓN Y TERMINACIÓN:")
    doc.add_paragraph(
        "El presente contrato tendrá una vigencia contada a partir de su firma hasta la finalización del evento y la correspondiente liquidación de los saldos pendientes. Podrá terminarse por mutuo acuerdo entre las partes o por incumplimiento grave de las obligaciones estipuladas."
    )

    # CLAUSULA DECIMA
    add_clause_heading("CLÁUSULA DÉCIMA –", "VALIDEZ Y FIRMAS ELECTRÓNICAS:")
    doc.add_paragraph(
        "De conformidad con la Ley 527 de 1999 de Colombia sobre Comercio Electrónico, el presente contrato mantendrá plena validez jurídica al ser firmado mediante firmas manuscritas escaneadas, firmas digitales o mediante el intercambio de mensajes de datos que demuestren la aceptación expresa de las partes."
    )

    doc.add_paragraph(
        "Para constancia de lo anterior, las partes firman el presente documento en dos (2) ejemplares del mismo tenor y validez, en la ciudad de [CIUDAD], el día [DÍA] del mes de [MES] del año 2026."
    )

    # Signature Block Table
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    cell_client = sig_table.rows[0].cells[0]
    cell_prov = sig_table.rows[0].cells[1]
    
    p_c = cell_client.paragraphs[0]
    p_c.add_run("_________________________________________\n").bold = True
    p_c.add_run("POR EL CONTRATANTE:\n").bold = True
    p_c.add_run("Nombre: [NOMBRE REPRESENTANTE]\nC.C. / NIT: [NÚMERO DOC]\nEmpresa: [NOMBRE CLIENTE]\nCargo: Representante Legal")

    p_p = cell_prov.paragraphs[0]
    p_p.add_run("_________________________________________\n").bold = True
    p_p.add_run("POR EL CONTRATISTA:\n").bold = True
    p_p.add_run("Nombre: [NOMBRE PROVEEDOR]\nC.C. / NIT: [NÚMERO DOC]\nEmpresa: CloudTickets\nCargo: Director / Líder Técnico")

    # Remove borders from signature table
    for row in sig_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders {} >
                    <w:top w:val="none"/>
                    <w:left w:val="none"/>
                    <w:bottom w:val="none"/>
                    <w:right w:val="none"/>
                </w:tcBorders>
            '''.format(nsdecls('w')))
            tcPr.append(tcBorders)

    # Save document
    doc.save(filename)
    print(f"Document saved successfully at: {filename}")

if __name__ == "__main__":
    create_contract_docx(r"c:\0DE\Ticketera\CONTRATO_SERVICIOS_CLOUDTICKETS_800000_COP.docx")
