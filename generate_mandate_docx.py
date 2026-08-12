import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_mandate_docx(filename):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Base Styles
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x27, 0x27, 0x2A)  # #27272a
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Palette
    NAVY = RGBColor(0x0F, 0x17, 0x2A)       # #0f172a
    DARK_BLUE = RGBColor(0x1E, 0x3A, 0x8A)  # #1e3a8a

    # Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CONTRATO DE MANDATO COMERCIAL PARA EL RECAUDO DE FONDOS,\nGESTORÍA DE BOLETERÍA Y SERVICIOS TECNOLÓGICOS")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = NAVY
    title_p.paragraph_format.space_after = Pt(4)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("(RECAUDO POR CUENTA DE TERCEROS - ARTÍCULO 1262 CÓDIGO DE COMERCIO)")
    sub_run.bold = True
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = DARK_BLUE
    subtitle_p.paragraph_format.space_after = Pt(18)

    # Preamble / Intro
    intro_p = doc.add_paragraph()
    intro_p.add_run("Entre los suscritos a saber: Por una parte, ")
    r_mandante = intro_p.add_run("[NOMBRE / RAZÓN SOCIAL DEL ORGANIZADOR]")
    r_mandante.bold = True
    intro_p.add_run(", persona jurídica/natural identificada con NIT / C.C. No. ")
    intro_p.add_run("[NIT O CÉDULA DEL MANDANTE]").bold = True
    intro_p.add_run(", domiciliada en la ciudad de [CIUDAD], representada legalmente por ")
    intro_p.add_run("[NOMBRE REPRESENTANTE LEGAL]").bold = True
    intro_p.add_run(", identificado(a) con C.C. No. [NÚMERO CÉDULA], quien en adelante y para todos los efectos del presente contrato se denominará ")
    intro_p.add_run("EL MANDANTE").bold = True
    intro_p.add_run("; y por la otra parte, ")
    r_mandatario = intro_p.add_run("[NOMBRE / RAZÓN SOCIAL PROVEEDOR CLOUDTICKETS]")
    r_mandatario.bold = True
    intro_p.add_run(", persona jurídica/natural identificada con NIT / C.C. No. ")
    intro_p.add_run("[NIT/C.C. DEL MANDATARIO]").bold = True
    intro_p.add_run(", domiciliada en [CIUDAD], quien opera la pasarela de pagos y plataforma digital CloudTickets, en adelante denominado ")
    intro_p.add_run("EL MANDATARIO").bold = True
    intro_p.add_run("; se ha convenido celebrar el presente ",)
    intro_p.add_run("CONTRATO DE MANDATO COMERCIAL PARA EL RECAUDO DE FONDOS POR CUENTA DE TERCEROS").bold = True
    intro_p.add_run(", sujeto a las siguientes cláusulas y disposiciones legales de la República de Colombia:")

    def add_clause_heading(number_title, text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        run1 = h.add_run(number_title)
        run1.bold = True
        run1.font.color.rgb = DARK_BLUE
        run1.font.size = Pt(11.5)
        run2 = h.add_run(" " + text)
        run2.bold = True
        run2.font.color.rgb = NAVY
        run2.font.size = Pt(11.5)
        return h

    # CLAUSULA PRIMERA
    add_clause_heading("CLÁUSULA PRIMERA –", "OBJETO DEL MANDATO Y ALCANCE:")
    doc.add_paragraph(
        "EL MANDANTE encomienda expresamente al MANDATARIO, y este acepta, la gestión comercial de recaudo por cuenta y riesgo del MANDANTE de los dineros provenientes de la venta y comercialización de boletería digital para el evento denominado:"
    )
    
    event_box = doc.add_paragraph()
    event_box.paragraph_format.left_indent = Inches(0.4)
    event_run = event_box.add_run(
        "• Nombre del Evento: [NOMBRE COMPLETO DEL EVENTO]\n"
        "• Fecha y Hora: [FECHA Y HORA DE REALIZACIÓN]\n"
        "• Lugar / Sede: [LUGAR / DIRECCIÓN / CIUDAD]\n"
        "• Aforo Máximo Autorizado: [CANTIDAD MÁXIMA DE BOLETAS A LA VENTA]"
    )
    event_run.bold = True
    event_run.font.color.rgb = DARK_BLUE

    doc.add_paragraph(
        "Para tal fin, EL MANDATARIO pondrá a disposición la plataforma tecnológica CloudTickets y su cuenta de pasarela de pagos Wompi Colombia y/o canales de transferencia bancaria autorizados."
    )

    # CLAUSULA SEGUNDA
    add_clause_heading("CLÁUSULA SEGUNDA –", "NATURALEZA DE LOS FONDOS Y EFECTOS FISCALES (ART. 1262 CÓDIGO DE COMERCIO):")
    doc.add_paragraph(
        "Las partes dejan constancia expresa de que el presente contrato se suscribe en los términos del Artículo 1262 y siguientes del Código de Comercio de Colombia. Por consiguiente:"
    )
    
    points_c2 = [
        ("1. Propiedad de los Fondos: ", "Los ingresos brutos recaudados por la venta de entradas pertenecen de manera exclusiva a EL MANDANTE desde el momento en que son abonados por el público comprador."),
        ("2. Recaudo para Terceros: ", "EL MANDATARIO actúa en calidad de simple mandatario recaudador por cuenta de terceros. Los dineros transferidos por los compradores a las cuentas o pasarelas del MANDATARIO NO constituyen ingreso computable, venta propia ni patrimonio del MANDATARIO ante la Dirección de Impuestos y Aduanas Nacionales (DIAN) ni entidades territoriales."),
        ("3. Ingreso Real del Mandatario: ", "Constituye ingreso propio únicamente el valor de la comisión por servicio y remuneración estipulada en la Cláusula Tercera del presente contrato.")
    ]
    for title, desc in points_c2:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(title).bold = True
        p.add_run(desc)

    # CLAUSULA TERCERA
    add_clause_heading("CLÁUSULA TERCERA –", "REMUNERACIÓN DEL MANDATARIO (COMISIONES):")
    doc.add_paragraph(
        "Como remuneración por la gestión del mandato, licenciamiento de la plataforma CloudTickets y soporte técnico, EL MANDANTE pagará al MANDATARIO los siguientes valores:"
    )

    fees = [
        ("A. Valor Fijo de Setup y Despliegue: ", "OCHOCIENTOS MIL PESOS COLOMBIANOS M/CTE ($800.000 COP), correspondientes a la configuración, infraestructura en la nube y puesta a punto de la plataforma."),
        ("B. Comisión Porcentual sobre Ventas: ", "El [PORCENTAJE %, ej. 4.0% ó 5.0%] sobre el valor bruto total de cada boleta vendida a través de la pasarela o plataforma."),
        ("C. Cargo Fijo por Ticket Emitido: ", "El valor de $[VALOR CARGO, ej. $350 COP] por cada boleta o entrada emitida exitosamente.")
    ]
    for title, desc in fees:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(title).bold = True
        p.add_run(desc)

    # CLAUSULA CUARTA
    add_clause_heading("CLÁUSULA CUARTA –", "DESCUENTOS OPERATIVOS, IMPUESTOS Y RETENCIONES:")
    doc.add_paragraph(
        "EL MANDATARIO queda expresa e irrevocablemente autorizado por EL MANDANTE para descontar directamente del valor bruto recaudado antes de proceder con el giro de saldos, los siguientes rubros operacionales:"
    )

    discounts = [
        ("1. Comisiones de Pasarela Wompi: ", "La tarifa cobrada por Wompi Colombia (2.9% + $590 COP + IVA por transacción exitosa o tarifa acordada)."),
        ("2. Retenciones Financieras: ", "Las retenciones en la fuente (ReteFuente, ReteIVA, ReteICA) practicadas automáticamente por la pasarela de pagos o entidades financieras sobre el recaudo."),
        ("3. Gravamen a los Movimientos Financieros (4x1000): ", "El valor equivalente al Impuesto del 4x1000 (GMF) generado en la cuenta bancaria del MANDATARIO al momento de realizar la transferencia de los fondos netos al MANDANTE."),
        ("4. Comisiones del Mandatario: ", "Los honorarios y comisiones acordados en la Cláusula Tercera.")
    ]
    for title, desc in discounts:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(title).bold = True
        p.add_run(desc)

    # CLAUSULA QUINTA
    add_clause_heading("CLÁUSULA QUINTA –", "ESQUEMA DE LIQUIDACIÓN, GIROS Y RESERVA DE GARANTÍA:")
    doc.add_paragraph(
        "La transferencia de los fondos netos recaudados a favor del MANDANTE se realizará bajo las siguientes condiciones de calendario y seguridad financiera:"
    )

    p_liq1 = doc.add_paragraph()
    p_liq1.paragraph_format.left_indent = Inches(0.2)
    p_liq1.add_run("A. Giros Parciales (Anticipos de Recaudo): ").bold = True
    p_liq1.add_run("EL MANDATARIO transferirá el [PORCENTAJE DE GIRO PARCIAL, ej. 80%] del recaudo neto acumulado con corte a la fecha [FECHA CORTE O DÍAS ANTES DEL EVENTO].")

    p_liq2 = doc.add_paragraph()
    p_liq2.paragraph_format.left_indent = Inches(0.2)
    p_liq2.add_run("B. Fondo de Reserva de Garantía: ").bold = True
    p_liq2.add_run("EL MANDATARIO retendrá un [PORCENTAJE DE RESERVA, ej. 20%] del valor neto total recaudado como fondo de reserva de garantía para cubrir eventuales solicitudes de devolución, retractos de compra o contracargos bancarios formulados por los compradores conforme a la Ley 1480 de 2011 (Estatuto del Consumidor).")

    p_liq3 = doc.add_paragraph()
    p_liq3.paragraph_format.left_indent = Inches(0.2)
    p_liq3.add_run("C. Liquidación Final y Desembolso del Saldo: ").bold = True
    p_liq3.add_run("Dentro de los tres a cinco (3 a 5) días hábiles siguientes a la finalización del evento, EL MANDATARIO enviará el estado de cuenta final con el desglose de ventas, deducciones e impuestos, y procederá a girar el saldo remanente a la cuenta bancaria del MANDANTE.")

    # CLAUSULA SEXTA
    add_clause_heading("CLÁUSULA SEXTA –", "RESPONSABILIDAD EXCLUSIVA SOBRE EL EVENTO Y DEVOLUCIONES:")
    doc.add_paragraph(
        "EL MANDANTE asume de forma exclusiva e ilimitada la responsabilidad civil, penal, comercial y administrativa frente a los compradores de entradas y autoridades públicas (Superintendencia de Industria y Comercio - SIC, Sayco & Acinpro, Alcaldías, etc.) por la realización, logística, calidad, aplazamiento o cancelación del evento."
    )
    doc.add_paragraph(
        "PARÁGRAFO PRIMERO: En caso de cancelación o aplazamiento del evento, EL MANDANTE autoriza al MANDATARIO a proceder con la devolución del dinero a los compradores utilizando los fondos recaudados o la reserva de garantía. Si los fondos en poder del MANDATARIO resultaren insuficientes para cubrir el 100% de los rembolsos, EL MANDANTE se obliga a transferir de forma inmediata la diferencia faltante al MANDATARIO en un plazo no mayor a cuarenta y ocho (48) horas."
    )

    # CLAUSULA SEPTIMA
    add_clause_heading("CLÁUSULA SÉPTIMA –", "OBLIGACIÓN DE FACTURACIÓN ANTE LA DIAN:")
    doc.add_paragraph(
        "En virtud de la naturaleza del mandato de recaudo, las obligaciones de facturación se distribuyen así:"
    )
    doc.add_paragraph().add_run("• Por parte del MANDATARIO: ").bold = True
    doc.paragraphs[-1].add_run("Emitirá Factura Electrónica de Venta o Cuenta de Cobro al MANDANTE ÚNICAMENTE por el valor de sus comisiones, honorarios de plataforma y servicio de despliegue.")
    
    doc.add_paragraph().add_run("• Por parte del MANDANTE: ").bold = True
    doc.paragraphs[-1].add_run("Es el único responsable de expedir la factura o documento equivalente de boletería al público comprador por el valor total de las entradas vendidas, así como de cumplir con sus obligaciones tributarias de Renta, IVA o Impuesto al Consumo según su régimen fiscal.")

    # CLAUSULA OCTAVA
    add_clause_heading("CLÁUSULA OCTAVA –", "CERTIFICACIÓN DE CUENTA BANCARIA DE DESTINO:")
    doc.add_paragraph(
        "Las transferencias y giros netos del recaudo se realizarán exclusivamente a la cuenta bancaria de la cual EL MANDANTE sea titular directo:"
    )

    bank_box = doc.add_paragraph()
    bank_box.paragraph_format.left_indent = Inches(0.4)
    bank_run = bank_box.add_run(
        "• Entidad Bancaria: [NOMBRE DEL BANCO]\n"
        "• Tipo de Cuenta: [AHORROS / CORRIENTE]\n"
        "• Número de Cuenta: [NÚMERO DE CUENTA]\n"
        "• Titular de la Cuenta: [NOMBRE / RAZÓN SOCIAL DEL TITULAR]\n"
        "• NIT / C.C. del Titular: [NIT O CÉDULA DE LA CUENTA]"
    )
    bank_run.bold = True
    bank_run.font.color.rgb = DARK_BLUE

    # CLAUSULA NOVENA
    add_clause_heading("CLÁUSULA NOVENA –", "DURACIÓN Y LEGISLACIÓN APLICABLE:")
    doc.add_paragraph(
        "El presente contrato rige a partir de la fecha de su firma y se extenderá hasta la liquidación final de saldos del evento. Todo lo no estipulado se regirá por el Código de Comercio de Colombia y la Ley 527 de 1999 de Comercio Electrónico para la validez de firmas digitales o escaneadas."
    )

    doc.add_paragraph(
        "Para constancia, las partes firman el presente documento en dos (2) ejemplares del mismo tenor y valor legal en la ciudad de [CIUDAD], el día [DÍA] del mes de [MES] del año 2026."
    )

    # Signatures
    doc.add_paragraph().paragraph_format.space_before = Pt(24)
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    cell_mandante = sig_table.rows[0].cells[0]
    cell_mandatario = sig_table.rows[0].cells[1]
    
    p_m = cell_mandante.paragraphs[0]
    p_m.add_run("_________________________________________\n").bold = True
    p_m.add_run("EL MANDANTE (ORGANIZADOR):\n").bold = True
    p_m.add_run("Firma: _________________________________\n"
              "Nombre: [NOMBRE REPRESENTANTE]\n"
              "C.C. / NIT: [NÚMERO DOC]\n"
              "Empresa: [NOMBRE ORGANIZADOR]")

    p_a = cell_mandatario.paragraphs[0]
    p_a.add_run("_________________________________________\n").bold = True
    p_a.add_run("EL MANDATARIO (CLOUDTICKETS / WOMPI):\n").bold = True
    p_a.add_run("Firma: _________________________________\n"
              "Nombre: [NOMBRE PROVEEDOR]\n"
              "C.C. / NIT: [NÚMERO DOC]\n"
              "Empresa: CloudTickets")

    # Remove borders from signature table
    for row in sig_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'''
                <w:tcBorders {} >
                    <w:top w:val="none"/>
                    <w:left w:val="none"/>
                    <w:bottom w:val="none"/>
                    <w:right w:val="none"/>
                </w:tcBorders>
            '''.format(nsdecls('w')))
            tcPr.append(tcBorders)

    doc.save(filename)
    print(f"Mandate contract DOCX saved successfully at: {filename}")

if __name__ == "__main__":
    create_mandate_docx(r"c:\0DE\Ticketera\CONTRATO_MANDATO_RECAUDO_CLOUDTICKETS.docx")
